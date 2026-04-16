"""문서 텍스트 추출 파이프라인 (v0.9).

지원 포맷:
- PDF: PyPDF2/pdfplumber (optional)
- DOCX: python-docx (optional)
- HTML: BeautifulSoup (optional)
- Markdown / 텍스트: 바로 파싱
- 표/그림 감지: 구조 메타데이터로 보존

외부 의존성 부재 시 graceful degradation — 텍스트만 추출.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from typing import Any

logger = logging.getLogger(__name__)


# === Data models ===


@dataclass
class ExtractedBlock:
    """추출된 문서 블록 (헤딩/문단/표/그림)."""
    type: str  # "heading" | "paragraph" | "table" | "image" | "list"
    content: str
    level: int | None = None  # heading level
    metadata: dict = field(default_factory=dict)


@dataclass
class ExtractionResult:
    """문서 추출 결과."""
    text: str  # 전체 평문
    blocks: list[ExtractedBlock]  # 구조화된 블록
    total_pages: int | None = None
    total_chars: int = 0
    detected_tables: int = 0
    detected_images: int = 0
    format: str = "unknown"
    extraction_errors: list[str] = field(default_factory=list)


# === Public API ===


def extract_from_bytes(
    content: bytes, filename: str | None = None, content_type: str | None = None,
) -> ExtractionResult:
    """파일 바이트에서 구조화된 텍스트 추출.

    Parameters
    ----------
    content : bytes
        파일 바이너리
    filename : str | None
        확장자 힌트용 (e.g., "lecture.pdf")
    content_type : str | None
        MIME 타입 힌트 (e.g., "application/pdf")

    Returns
    -------
    ExtractionResult
    """
    fmt = _detect_format(filename, content_type, content)
    logger.info("Extracting document: format=%s size=%d", fmt, len(content))

    try:
        if fmt == "pdf":
            return _extract_pdf(content)
        elif fmt == "docx":
            return _extract_docx(content)
        elif fmt == "html":
            return _extract_html(content)
        elif fmt in ("markdown", "text"):
            return _extract_text(content, fmt)
        else:
            # 알 수 없는 형식 → 텍스트로 시도
            return _extract_text(content, "unknown")
    except Exception as exc:  # noqa: BLE001
        logger.exception("Extraction failed: %s", exc)
        return ExtractionResult(
            text="",
            blocks=[],
            format=fmt,
            extraction_errors=[repr(exc)],
        )


def extract_from_text(text: str) -> ExtractionResult:
    """플레인 텍스트 / 마크다운에서 블록 구조 추출."""
    blocks = _parse_markdown_blocks(text)
    return ExtractionResult(
        text=text,
        blocks=blocks,
        total_chars=len(text),
        detected_tables=sum(1 for b in blocks if b.type == "table"),
        detected_images=sum(1 for b in blocks if b.type == "image"),
        format="markdown",
    )


# === Format detection ===


def _detect_format(
    filename: str | None, content_type: str | None, content: bytes,
) -> str:
    """파일 형식 탐지 (확장자 → MIME → 매직 바이트 순)."""
    if filename:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext == "pdf":
            return "pdf"
        if ext in ("docx", "doc"):
            return "docx"
        if ext in ("html", "htm"):
            return "html"
        if ext in ("md", "markdown"):
            return "markdown"
        if ext in ("txt",):
            return "text"

    if content_type:
        ct = content_type.lower()
        if "pdf" in ct:
            return "pdf"
        if "wordprocessingml" in ct or "msword" in ct:
            return "docx"
        if "html" in ct:
            return "html"
        if "markdown" in ct:
            return "markdown"
        if "text/plain" in ct:
            return "text"

    # 매직 바이트
    if content[:4] == b"%PDF":
        return "pdf"
    if content[:2] == b"PK":  # ZIP 시그니처 (DOCX/XLSX)
        return "docx"

    return "text"


# === PDF extraction ===


def _extract_pdf(content: bytes) -> ExtractionResult:
    """PDF 텍스트 추출 — pdfplumber 우선, PyPDF2 fallback."""
    blocks: list[ExtractedBlock] = []
    text_parts: list[str] = []
    total_pages = 0
    tables_count = 0
    errors: list[str] = []

    # pdfplumber 시도 (표 추출 우수)
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            total_pages = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
                    blocks.append(ExtractedBlock(
                        type="paragraph",
                        content=page_text,
                        metadata={"page": i + 1},
                    ))
                # 표 추출
                try:
                    tables = page.extract_tables()
                    for t_idx, table in enumerate(tables or []):
                        tables_count += 1
                        table_text = "\n".join(
                            " | ".join(str(cell or "") for cell in row) for row in table
                        )
                        blocks.append(ExtractedBlock(
                            type="table",
                            content=table_text,
                            metadata={"page": i + 1, "table_index": t_idx, "rows": len(table)},
                        ))
                        text_parts.append(f"\n[표 {tables_count}]\n{table_text}\n")
                except Exception as exc:  # noqa: BLE001
                    errors.append(f"table extraction failed on page {i+1}: {exc}")
    except ImportError:
        # PyPDF2 fallback
        try:
            import PyPDF2  # type: ignore
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            total_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
                    blocks.append(ExtractedBlock(
                        type="paragraph",
                        content=page_text,
                        metadata={"page": i + 1},
                    ))
        except ImportError:
            errors.append("Neither pdfplumber nor PyPDF2 is installed")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"PDF parsing failed: {exc}")

    full_text = "\n\n".join(text_parts)
    return ExtractionResult(
        text=full_text,
        blocks=blocks,
        total_pages=total_pages,
        total_chars=len(full_text),
        detected_tables=tables_count,
        detected_images=0,  # TODO: pdfplumber images
        format="pdf",
        extraction_errors=errors,
    )


# === DOCX extraction ===


def _extract_docx(content: bytes) -> ExtractionResult:
    """DOCX 텍스트 추출 — python-docx 사용."""
    blocks: list[ExtractedBlock] = []
    text_parts: list[str] = []
    tables_count = 0
    errors: list[str] = []

    try:
        from docx import Document  # type: ignore
        doc = Document(io.BytesIO(content))

        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                level = int(style.replace("Heading ", "") or 1) if "Heading " in style else 1
                blocks.append(ExtractedBlock(
                    type="heading", content=txt, level=level,
                ))
            else:
                blocks.append(ExtractedBlock(type="paragraph", content=txt))
            text_parts.append(txt)

        for i, table in enumerate(doc.tables):
            tables_count += 1
            table_text = "\n".join(
                " | ".join(cell.text for cell in row.cells) for row in table.rows
            )
            blocks.append(ExtractedBlock(
                type="table", content=table_text,
                metadata={"table_index": i, "rows": len(table.rows)},
            ))
            text_parts.append(f"\n[표 {tables_count}]\n{table_text}\n")
    except ImportError:
        errors.append("python-docx is not installed")
    except Exception as exc:  # noqa: BLE001
        errors.append(f"DOCX parsing failed: {exc}")

    full_text = "\n\n".join(text_parts)
    return ExtractionResult(
        text=full_text,
        blocks=blocks,
        total_chars=len(full_text),
        detected_tables=tables_count,
        format="docx",
        extraction_errors=errors,
    )


# === HTML extraction ===


def _extract_html(content: bytes) -> ExtractionResult:
    """HTML 텍스트 추출 — BeautifulSoup 사용."""
    blocks: list[ExtractedBlock] = []
    text_parts: list[str] = []
    tables_count = 0
    images_count = 0
    errors: list[str] = []

    try:
        from bs4 import BeautifulSoup  # type: ignore
        soup = BeautifulSoup(content, "html.parser")

        # 스크립트/스타일 제거
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()

        # 헤딩
        for level in range(1, 7):
            for h in soup.find_all(f"h{level}"):
                txt = h.get_text(strip=True)
                if txt:
                    blocks.append(ExtractedBlock(
                        type="heading", content=txt, level=level,
                    ))

        # 문단
        for p in soup.find_all(["p", "li"]):
            txt = p.get_text(strip=True)
            if txt:
                blocks.append(ExtractedBlock(
                    type="list" if p.name == "li" else "paragraph",
                    content=txt,
                ))
                text_parts.append(txt)

        # 표
        for i, table in enumerate(soup.find_all("table")):
            tables_count += 1
            rows = []
            for row in table.find_all("tr"):
                cells = [c.get_text(strip=True) for c in row.find_all(["td", "th"])]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            blocks.append(ExtractedBlock(
                type="table", content=table_text,
                metadata={"table_index": i, "rows": len(rows)},
            ))
            text_parts.append(f"\n[표 {tables_count}]\n{table_text}\n")

        # 이미지 (alt 텍스트만 보존)
        for img in soup.find_all("img"):
            alt = img.get("alt", "")
            if alt:
                images_count += 1
                blocks.append(ExtractedBlock(
                    type="image", content=f"[이미지: {alt}]",
                    metadata={"src": img.get("src"), "alt": alt},
                ))
    except ImportError:
        # BeautifulSoup 없으면 정규식 fallback
        text = content.decode("utf-8", errors="replace")
        text = re.sub(r"<[^>]+>", "", text)
        text = re.sub(r"\s+", " ", text).strip()
        text_parts.append(text)
        blocks.append(ExtractedBlock(type="paragraph", content=text))
    except Exception as exc:  # noqa: BLE001
        errors.append(f"HTML parsing failed: {exc}")

    full_text = "\n\n".join(text_parts)
    return ExtractionResult(
        text=full_text,
        blocks=blocks,
        total_chars=len(full_text),
        detected_tables=tables_count,
        detected_images=images_count,
        format="html",
        extraction_errors=errors,
    )


# === Text / Markdown ===


def _extract_text(content: bytes, fmt: str) -> ExtractionResult:
    """텍스트/마크다운 디코딩 + 블록 파싱."""
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("utf-8", errors="replace")

    blocks = _parse_markdown_blocks(text) if fmt == "markdown" else [
        ExtractedBlock(type="paragraph", content=text)
    ]
    return ExtractionResult(
        text=text,
        blocks=blocks,
        total_chars=len(text),
        format=fmt,
    )


def _parse_markdown_blocks(text: str) -> list[ExtractedBlock]:
    """간단한 마크다운 블록 파서."""
    blocks: list[ExtractedBlock] = []
    lines = text.split("\n")
    buffer: list[str] = []

    def flush_para():
        if buffer:
            content = "\n".join(buffer).strip()
            if content:
                blocks.append(ExtractedBlock(type="paragraph", content=content))
            buffer.clear()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 헤딩
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            flush_para()
            level = len(m.group(1))
            blocks.append(ExtractedBlock(type="heading", content=m.group(2).strip(), level=level))
            i += 1
            continue

        # 리스트
        if re.match(r"^\s*[-*+]\s+", line) or re.match(r"^\s*\d+\.\s+", line):
            flush_para()
            list_items = []
            while i < len(lines) and (
                re.match(r"^\s*[-*+]\s+", lines[i]) or re.match(r"^\s*\d+\.\s+", lines[i])
            ):
                list_items.append(lines[i])
                i += 1
            blocks.append(ExtractedBlock(type="list", content="\n".join(list_items)))
            continue

        # 표 (간단한 pipe table)
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?[-:\s|]+\|?\s*$", lines[i + 1]):
            flush_para()
            table_lines = [line]
            i += 1  # separator
            i += 1
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            blocks.append(ExtractedBlock(
                type="table", content="\n".join(table_lines),
                metadata={"rows": len(table_lines)},
            ))
            continue

        # 이미지 markdown
        m_img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if m_img:
            flush_para()
            alt, src = m_img.groups()
            blocks.append(ExtractedBlock(
                type="image", content=f"[이미지: {alt}]",
                metadata={"src": src, "alt": alt},
            ))
            i += 1
            continue

        # 빈 줄 → 문단 경계
        if not stripped:
            flush_para()
        else:
            buffer.append(line)
        i += 1

    flush_para()
    return blocks


# === Metadata tagging ===


def auto_tag_content(text: str, department: str | None = None) -> list[str]:
    """자동 태그 추출 — 키워드 기반 휴리스틱.

    의료/간호 도메인 핵심어가 나타나면 태그로 추가.
    """
    tags: set[str] = set()
    text_lower = text.lower()

    # 도메인 키워드
    keywords = {
        "간호": ["간호", "nursing", "투약", "활력징후", "무균"],
        "해부학": ["해부", "anatomy", "근골격", "신경계", "순환계"],
        "약리학": ["약리", "pharmacology", "약물", "용량", "부작용"],
        "병리학": ["병리", "pathology", "질환", "염증", "종양"],
        "기초의학": ["생리", "생화학", "미생물", "면역"],
        "물리치료": ["물리치료", "운동치료", "근골격계", "재활"],
        "치위생": ["구강", "치주", "스케일링", "치과"],
    }

    for tag, kw_list in keywords.items():
        if any(kw in text_lower for kw in kw_list):
            tags.add(tag)

    # 국가고시 관련
    if any(kw in text_lower for kw in ["국가고시", "국시", "national exam"]):
        tags.add("국가고시")

    # 학과 태그
    if department:
        tags.add(department.lower())

    return sorted(tags)
